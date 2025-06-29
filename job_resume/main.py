import os
import tempfile
import uuid
from flask import Flask, render_template, request, jsonify, send_file
from werkzeug.utils import secure_filename
import PyPDF2
import docx
from datetime import datetime
from io import BytesIO
from langchain_openai import ChatOpenAI
from langchain.prompts import PromptTemplate
from langchain_core.runnables.base import RunnableSequence
from langchain.output_parsers import PydanticOutputParser
from pydantic import BaseModel, Field
from typing import List, Dict
import json
from docx import Document as DocxDocument
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from dotenv import load_dotenv

# Initialize Flask app
app = Flask(__name__, static_folder='static')
app.config['UPLOAD_FOLDER'] = tempfile.mkdtemp()
app.config['MAX_CONTENT_LENGTH'] = 5 * 1024 * 1024  # 5MB max file size

# Configure environment variables
load_dotenv()
openai_api_key = os.getenv("OPENAI_API_KEY")

# Initialize language model
llm = ChatOpenAI(
    model_name="gpt-4o-mini",
    temperature=0.2,
    openai_api_key=openai_api_key
)

# Define Pydantic models for structured output parsing
class ATSScore(BaseModel):
    keyword_match_percentage: float = Field(description="Percentage of job description keywords found in resume")
    keyword_frequency_score: float = Field(description="Score based on how often keywords appear in the resume")
    section_completion_percentage: float = Field(description="Percentage of key sections present in the resume")
    formatting_readability_score: float = Field(description="Score for overall formatting and readability")
    hard_soft_skills_balance: float = Field(description="Score for balance between hard technical and soft interpersonal skills")
    proximity_score: float = Field(description="Score for phrases appearing in the right order/context")
    total_ats_score: float = Field(description="Final weighted ATS compatibility score out of 100")
    missing_keywords: List[str] = Field(description="Important keywords from job description missing in resume")
    improvement_suggestions: List[str] = Field(description="Specific suggestions to improve ATS compatibility")
    
    # New fields for categorized suggestions
    searchability_suggestions: List[str] = Field(description="Suggestions to improve keyword searchability")
    skills_suggestions: List[str] = Field(description="Suggestions to improve skills presentation")
    formatting_suggestions: List[str] = Field(description="Suggestions to improve formatting and readability")
    section_suggestions: List[str] = Field(description="Suggestions to improve section completeness")
    synonym_suggestions: List[str] = Field(description="Suggestions to improve keyword variation coverage")

    # issue‐count fields:
    searchability_issues_count: int = Field(
        description="Number of searchability issues to fix"
    )
    skills_issues_count: int = Field(
        description="Number of skills‐presentation issues to fix"
    )
    formatting_issues_count: int = Field(
        description="Number of formatting/readability issues to fix"
    )
    section_issues_count: int = Field(
        description="Number of section completeness issues to fix"
    )
    synonym_issues_count: int = Field(
        description="Number of keyword‐variation issues to fix"
    )

class ResumeOptimization(BaseModel):
    improved_summary: str = Field(description="AI-enhanced professional summary")
    improved_bullets: Dict[str, List[str]] = Field(description="Improved bullet points for each experience section")
    suggested_skills: List[str] = Field(description="Additional skills to highlight based on job description")
    formatting_suggestions: List[str] = Field(description="Suggestions for better formatting")
    improved_resume_text: str = Field(description="Full improved resume text")

class CoverLetterOutput(BaseModel):
    cover_letter_text: str = Field(description="Complete cover letter text")

# Helper functions for file processing
def extract_text_from_pdf(file_path):
    text = ""
    with open(file_path, 'rb') as file:
        pdf_reader = PyPDF2.PdfReader(file)
        for page in pdf_reader.pages:
            text += page.extract_text()
    return text

def extract_text_from_docx(file_path):
    doc = docx.Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

def process_resume_file(file):
    filename = secure_filename(file.filename)
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file.save(file_path)
    
    if filename.endswith('.pdf'):
        resume_text = extract_text_from_pdf(file_path)
    elif filename.endswith('.docx') or filename.endswith('.doc'):
        resume_text = extract_text_from_docx(file_path)
    else:
        return None, "Unsupported file format"
    
    return resume_text, file_path

# ATS Score Analysis Chain
def create_ats_analysis_chain():
    ats_template = """
    You are an expert ATS (Applicant Tracking System) analyzer and resume optimization specialist.

    I need you to analyze a resume against a specific job description and provide a detailed ATS compatibility score.

    RESUME TEXT:
    {resume_text}

    JOB DESCRIPTION:
    {job_description}

   Analyze the following resume against the provided job description and generate a detailed ATS compatibility report including:

    1. Keyword Match %:
    - Extract ALL critical keywords and key phrases (skills, certifications, tools, industry jargon, job titles).
    - Assess exact and semantic keyword matches, including synonyms and relevant variations.
    - Evaluate keyword placement priority: prioritize headline, professional summary, skills section, and work experience.
    - Score keyword density ensuring natural language flow (avoid stuffing).
    - Identify missing high-impact keywords critical to the job role and recommend optimal placement.

    2. Section Completion %:
    - Verify presence of core ATS sections with professional formatting.
    - Score based on alignment with job requirements and ATS parsing ease.
    - Recommend missing or reorder sections for maximum ATS compatibility.

    3. Formatting and Readability Score:
    Assess resume formatting and readability for ATS systems, including:
    - Use of standard fonts (e.g., Arial, Calibri)
    - Consistent date formats (e.g., MM/YYYY)
    - Absence of graphics, tables, columns, headers, or footers
    - Clear section headings and logical structure
    - Appropriate file format (.docx or .pdf)

    4. Hard vs Soft Skills Balance:
    Evaluate the balance between technical (hard) and interpersonal (soft) skills aligned to the job description.

    5. Proximity Score:
    Analyze whether related keywords and phrases appear close together and in the correct context.

    6. Final ATS Score:
    Calculate a weighted ATS score out of 100 using the formula:
    (Keyword Match * 0.35) + (Section Completion * 0.25) + (Formatting * 0.20) + (Skills Balance * 0.10) + (Proximity * 0.10)

    7. List important keywords from the job description missing in the resume

    Additionally, organize your suggestions into the following categories to match with the job description and get a better ATS score:
    1. Searchability suggestions - How to improve keyword match and searchability
       • Also provide `searchability_issues_count`: the number of searchability issues to fix
    2. Skills suggestions - How to improve hard/soft skills balance and presentation
       • Also provide `skills_issues_count`: the number of skills-presentation issues to fix
    3. Formatting suggestions - How to improve layout and ATS readability
       • Also provide `formatting_issues_count`: the number of formatting/readability issues to fix
    4. Section suggestions - How to improve section completeness
       • Also provide `section_issues_count`: the number of section completeness issues to fix
    5. Synonym suggestions - How to improve keyword variation coverage
       • Also provide `synonym_issues_count`: the number of keyword-variation issues to fix

    For each category:
    1. List all individual issues that need fixing.
    2. Provide 2–3 specific, actionable suggestions.
    3. Emit the corresponding `<category>_issues_count` integer, **exactly equal** to the number of issues you listed above.

    {format_instructions}
    """
    
    parser = PydanticOutputParser(pydantic_object=ATSScore)
    prompt = PromptTemplate(
        template=ats_template,
        input_variables=["resume_text", "job_description"],
        partial_variables={"format_instructions": parser.get_format_instructions()}
    )

    chain = RunnableSequence(prompt, llm, parser)
    return chain

# Resume Optimization Chain
def create_resume_optimization_chain():
    resume_template = """
    You are an expert resume writer, career coach, and Applicant Tracking System (ATS) specialist. Your mission is to transform the candidate’s existing resume into a highly optimized, keyword‑rich document that perfectly aligns with the given job description—while **preserving every original section** of the resume.

    **INPUTS:**  
    - **RESUME TEXT:** `{resume_text}`  
    - **JOB DESCRIPTION:** `{job_description}`  
    - **ATS ANALYSIS:** `{ats_analysis}`  

    ---

    ## Instructions

    1. **Section Preservation & Enhancement**  
    - **Retain every section** from the original resume (e.g., Professional Summary, Experience, Projects, Education, Skills, Certifications, Volunteer, etc.).  
    - For each section, **refine content** to mirror the job description’s language and keywords.

    2. **Skill Extraction & Mapping**  
    - Parse the JD for all **required** and **preferred** skills/technologies/methodologies.  
    - Build a “Skills Checklist” and ensure each skill appears in the resume—either in the Skills section or embedded within relevant bullets.  
    - If the candidate legitimately holds a JD‑listed skill that’s missing, add it; do not invent.

    3. **Professional Summary (1–3 sentences)**  
    - Incorporate the top 4–6 JD keywords/phrases.  
    - Highlight years of experience, core qualifications, and career goal.  
    - Keep it concise, impactful, and natural.

    4. **Experience Section Optimization**  
    - Maintain **reverse‑chronological** order and use present‑tense for current roles, past‑tense for prior roles.  
    - **Rewrite each bullet** to:  
        - Use strong action verbs (e.g., “Led,” “Architected,” “Streamlined”).  
        - Embed exact JD keywords (e.g., “CI/CD pipelines,” “Agile,” “AWS”).  
        - Quantify achievements (e.g., “Reduced deployment time by 40%,” “Managed a $200K budget”).  
    - **Retain all original positions** and headers; enhance every bullet for relevance and metrics.

    5. **Skills Section Enhancement**  
    - Under a clear **Skills** heading, list every core JD skill/term as individual bullets, ordered by relevance.  
    - If a required skill is absent in the original, note the gap and suggest a related, transferable skill.

    6. **Education & Certifications**  
    - **Preserve** all original entries.  
    - Standardize formatting: *Degree/Certification* – *Institution/Issuer*, *Month Year*.  
    - Add any JD‑required certification the candidate holds or is pursuing.

    7. **Formatting for Optimal ATS Parsing**  
    - Single‑column, reverse‑chronological layout with standard headings.  
    - Simple bullets (– or ●); no tables, graphics, or columns.  
    - Standard font (Arial or Calibri, 10–12 pt); consistent spacing, margins, and date format (e.g., “MMM YYYY – MMM YYYY”).  
    - Recommend final filename: `Firstname_Lastname_Resume.pdf`.

    8. **Full Resume Rewrite**  
    - Provide a polished 1–2 page document that:  
        1. **Professional Summary** (refined)  
        2. **Experience** (all original roles, optimized)  
        3. **Skills** (exhaustive, JD‑aligned)  
        4. **Education** (standardized)  
        5. **Certifications** (standardized)  
        6. **Other Sections** (e.g., Projects, Volunteer) – refined to include relevant keywords.  
    - Ensure **every JD keyword** appears at least once across summary, skills, or experience—without unnatural repetition.

    9. **Final Verification**  
    - Confirm 100% coverage of required/preferred JD skills.  
    - Ensure natural keyword density and no keyword stuffing.  
    - Validate that each original section is present and improved.  
    - Include formatting notes in curly braces (e.g., `{format_instructions}`) for the user’s final layout.

    **OUTPUT:**  
    A complete ATS‑optimized resume—retaining and enhancing all original sections.
    """
    
    parser = PydanticOutputParser(pydantic_object=ResumeOptimization)
    prompt = PromptTemplate(
        template=resume_template,
        input_variables=["resume_text", "job_description", "ats_analysis"],
        partial_variables={"format_instructions": parser.get_format_instructions()}
    )

    chain = RunnableSequence(prompt, llm, parser)
    return chain

# Cover Letter Generation Chain
def create_cover_letter_chain():
    current_date = datetime.now().strftime("%B %d, %Y")
    cover_letter_template = """
    You are an expert cover letter writer with deep knowledge of professional communication and hiring practices.
    
    Create a compelling, personalized cover letter based on the resume and job description provided.
    
    RESUME TEXT:
    {resume_text}
    
    JOB DESCRIPTION:
    {job_description}

    CURRENT DATE: {current_date}
    
    Please write a professional cover letter that:
    1. Has a proper business letter format with date and contact information
    2. Opens with a compelling introduction that shows enthusiasm for the position
    3. Highlights 2-3 key qualifications from the resume that directly match the job requirements
    4. Incorporates important keywords from the job description naturally
    5. Explains why the applicant is a good fit for the company and role specifically
    6. Closes with a strong call to action and thank you
    7. Maintains a professional yet personable tone
    8. Is approximately 250-350 words in total
    
    {format_instructions}
    """
    
    parser = PydanticOutputParser(pydantic_object=CoverLetterOutput)
    prompt = PromptTemplate(
        template=cover_letter_template,
        input_variables=["resume_text", "job_description"],
        partial_variables={"format_instructions": parser.get_format_instructions(),
                         "current_date": current_date}
    )
    
    return RunnableSequence(prompt, llm, parser)

# Document creation functions
def create_docx_document(content, document_type="resume"):
    doc = DocxDocument()
    
    # Add content
    sections = content.split('\n\n')
    for section in sections:
        if section.strip():
            doc.add_paragraph(section)
    
    # Save document to memory
    file_obj = BytesIO()
    doc.save(file_obj)
    file_obj.seek(0)
    
    return file_obj

def create_pdf_document(content, document_type="resume"):
    file_obj = BytesIO()
    
    # Create PDF
    doc = SimpleDocTemplate(file_obj, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    
    normal_style = styles['Normal']
    heading_style = styles['Heading1']
    
    # Parse content and add to document
    sections = content.split('\n\n')
    for section in sections:
        if section.strip():
            # Check if it's a heading (simple heuristic)
            if len(section) < 50 and section.isupper() or ':' in section and len(section.split(':')[0]) < 20:
                story.append(Paragraph(section, heading_style))
            else:
                story.append(Paragraph(section, normal_style))
            story.append(Spacer(1, 12))
    
    doc.build(story)
    file_obj.seek(0)
    return file_obj

# Flask routes - Updated for separate ATS analysis and cover letter generation
@app.route('/')
def index():
    return render_template('index.html')

@app.route('/analyze-ats', methods=['POST'])
def analyze_ats():
    """Analyze resume for ATS compatibility and generate optimized version"""
    if 'resume' not in request.files:
        return jsonify({'error': 'No resume file uploaded'}), 400
    
    resume_file = request.files['resume']
    job_description = request.form.get('job_description', '')
    
    if resume_file.filename == '':
        return jsonify({'error': 'No resume file selected'}), 400
    
    if not job_description:
        return jsonify({'error': 'Job description is required'}), 400
    
    # Process the resume file
    resume_text, file_path = process_resume_file(resume_file)
    if not resume_text:
        return jsonify({'error': 'Could not extract text from resume'}), 400
    
    # Run ATS analysis on original resume
    ats_chain = create_ats_analysis_chain()
    original_ats_result = ats_chain.invoke({"resume_text": resume_text, "job_description": job_description})

    # Run resume optimization
    optimization_chain = create_resume_optimization_chain()
    optimization_result = optimization_chain.invoke({
        "resume_text": resume_text,
        "job_description": job_description,
        "ats_analysis": json.dumps(original_ats_result.model_dump())
    })

    # Create a unique session ID for this analysis
    session_id = str(uuid.uuid4())
    
    # Store results in temporary files for later retrieval
    temp_dir = os.path.join(app.config['UPLOAD_FOLDER'], session_id)
    os.makedirs(temp_dir, exist_ok=True)
    
    # Store original resume text, job description
    with open(os.path.join(temp_dir, 'ats_data.json'), 'w') as f:
        json.dump({
            'resume_text': resume_text,
            'job_description': job_description,
            'original_ats_analysis': original_ats_result.model_dump(),
            'optimization_result': optimization_result.model_dump()
        }, f)
    
    # Return results with both original and optimized scores
    return jsonify({
        'session_id': session_id,
        'ats_analysis': original_ats_result.model_dump(),
        'optimization_result': optimization_result.model_dump()
    })

@app.route('/generate-cover-letter', methods=['POST'])
def generate_cover_letter():
    """Generate cover letter separately"""
    if 'resume' not in request.files:
        return jsonify({'error': 'No resume file uploaded'}), 400
    
    resume_file = request.files['resume']
    job_description = request.form.get('job_description', '')
    
    if resume_file.filename == '':
        return jsonify({'error': 'No resume file selected'}), 400
    
    if not job_description:
        return jsonify({'error': 'Job description is required'}), 400
    
    # Process the resume file
    resume_text, file_path = process_resume_file(resume_file)
    if not resume_text:
        return jsonify({'error': 'Could not extract text from resume'}), 400
    
    # Generate cover letter
    cover_letter_chain = create_cover_letter_chain()
    cover_letter_result = cover_letter_chain.invoke({
        "resume_text": resume_text,
        "job_description": job_description
    })
    
    # Create a unique session ID for this cover letter
    session_id = str(uuid.uuid4())
    
    # Store results in temporary files for later retrieval
    temp_dir = os.path.join(app.config['UPLOAD_FOLDER'], session_id)
    os.makedirs(temp_dir, exist_ok=True)
    
    # Store original data for regeneration
    with open(os.path.join(temp_dir, 'cover_letter_data.json'), 'w') as f:
        json.dump({
            'resume_text': resume_text,
            'job_description': job_description,
            'cover_letter': cover_letter_result.model_dump()  
        }, f)
    
    return jsonify({
        'session_id': session_id,
        'cover_letter': cover_letter_result.cover_letter_text  
    })

@app.route('/regenerate-ats/<session_id>', methods=['POST'])
def regenerate_ats(session_id):
    """Regenerate ATS analysis and optimization"""
    temp_dir = os.path.join(app.config['UPLOAD_FOLDER'], session_id)
    data_file = os.path.join(temp_dir, 'ats_data.json')
    
    if not os.path.exists(data_file):
        return jsonify({'error': 'Session data not found'}), 404
    
    with open(data_file, 'r') as f:
        data = json.load(f)
    
    resume_text = data['resume_text']
    job_description = data['job_description']
    
    # Re-run ATS analysis on original resume
    ats_chain = create_ats_analysis_chain()
    original_ats_result = ats_chain.invoke({
        "resume_text": resume_text,
        "job_description": job_description
    })
    
    # Re-run resume optimization
    optimization_chain = create_resume_optimization_chain()
    optimization_result = optimization_chain.invoke({
        "resume_text": resume_text,
        "job_description": job_description,
        "ats_analysis": json.dumps(original_ats_result.model_dump())
    })
    
    # Update stored data
    with open(data_file, 'w') as f:
        json.dump({
            'resume_text': resume_text,
            'job_description': job_description,
            'original_ats_analysis': original_ats_result.model_dump(),
            'optimization_result': optimization_result.model_dump()
        }, f)
    
    return jsonify({
        'ats_analysis': original_ats_result.model_dump(),
        'optimization_result': optimization_result.model_dump()
    })

@app.route('/regenerate-cover-letter/<session_id>', methods=['POST'])
def regenerate_cover_letter(session_id):
    """Regenerate cover letter"""
    temp_dir = os.path.join(app.config['UPLOAD_FOLDER'], session_id)
    data_file = os.path.join(temp_dir, 'cover_letter_data.json')
    
    if not os.path.exists(data_file):
        return jsonify({'error': 'Session data not found'}), 404
    
    with open(data_file, 'r') as f:
        data = json.load(f)
    
    resume_text = data['resume_text']
    job_description = data['job_description']
    
    # Re-generate cover letter
    cover_letter_chain = create_cover_letter_chain()
    cover_letter_result = cover_letter_chain.invoke({
        "resume_text": resume_text,
        "job_description": job_description
    })
    
    # Update stored data 
    with open(data_file, 'w') as f:
        json.dump({
            'resume_text': resume_text,
            'job_description': job_description,
            'cover_letter': cover_letter_result.model_dump()  
        }, f)
    
    return jsonify({
        'cover_letter': cover_letter_result.cover_letter_text  
    })

@app.route('/preview/<document_type>/<session_id>')
def preview_document(document_type, session_id):
    """Preview document before download with score comparison"""
    temp_dir = os.path.join(app.config['UPLOAD_FOLDER'], session_id)
    
    if document_type == 'resume':
        data_file = os.path.join(temp_dir, 'ats_data.json')
        if not os.path.exists(data_file):
            return jsonify({'error': 'Session data not found'}), 404
        
        with open(data_file, 'r') as f:
            data = json.load(f)
        
        # Run optimized ATS analysis if not already done
        if 'optimized_ats_analysis' not in data:
            ats_chain = create_ats_analysis_chain()
            optimized_ats_result = ats_chain.invoke({
                "resume_text": data['optimization_result']['improved_resume_text'],
                "job_description": data['job_description']
            })
            data['optimized_ats_analysis'] = optimized_ats_result.model_dump()
            with open(data_file, 'w') as f:
                json.dump(data, f)
        
        content = data['optimization_result']['improved_resume_text']
        original_score = data['original_ats_analysis']['total_ats_score']
        optimized_score = data.get('optimized_ats_analysis', {}).get('total_ats_score', 0)
        
        return jsonify({
            'content': content,
            'score_comparison': {
                'original_score': original_score,
                'optimized_score': optimized_score
            }
        })
        
    elif document_type == 'cover_letter':
        data_file = os.path.join(temp_dir, 'cover_letter_data.json')
        if not os.path.exists(data_file):
            return jsonify({'error': 'Session data not found'}), 404
        
        with open(data_file, 'r') as f:
            data = json.load(f)
        
        content = data['cover_letter']['cover_letter_text']
        return jsonify({'content': content})
    else:
        return jsonify({'error': 'Invalid document type'}), 400

@app.route('/download/<file_type>/<document_type>/<session_id>')
def download_document(file_type, document_type, session_id):
    """Download document in specified format"""
    temp_dir = os.path.join(app.config['UPLOAD_FOLDER'], session_id)
    
    if document_type == 'resume':
        data_file = os.path.join(temp_dir, 'ats_data.json')
        if not os.path.exists(data_file):
            return jsonify({'error': 'Session data not found'}), 404
        
        with open(data_file, 'r') as f:
            data = json.load(f)
        
        content = data['optimization_result']['improved_resume_text']
        filename = f"optimized_resume.{file_type}"
        
    elif document_type == 'cover_letter':
        data_file = os.path.join(temp_dir, 'cover_letter_data.json')
        if not os.path.exists(data_file):
            return jsonify({'error': 'Session data not found'}), 404
        
        with open(data_file, 'r') as f:
            data = json.load(f)
        
        content = data['cover_letter']['cover_letter_text']
        filename = f"cover_letter.{file_type}"
    else:
        return jsonify({'error': 'Invalid document type'}), 400
    
    # Generate requested file format
    if file_type == 'docx':
        file_obj = create_docx_document(content, document_type)
        mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    elif file_type == 'pdf':
        file_obj = create_pdf_document(content, document_type)
        mimetype = 'application/pdf'
    else:
        return jsonify({'error': 'Invalid file type'}), 400
    
    return send_file(
        file_obj,
        mimetype=mimetype,
        as_attachment=True,
        download_name=filename
    )

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)