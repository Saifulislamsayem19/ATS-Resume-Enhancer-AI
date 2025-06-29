import os
import asyncio
import PyPDF2
import docx
from datetime import datetime
from io import BytesIO
from werkzeug.utils import secure_filename
from langchain_openai import ChatOpenAI
from langchain.prompts import PromptTemplate
from langchain_core.runnables.base import RunnableSequence
from langchain.output_parsers import PydanticOutputParser
from pydantic import BaseModel, Field
from typing import List, Dict
from docx import Document as DocxDocument
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

# Define Pydantic models
class ATSScore(BaseModel):
    keyword_match_percentage: float = Field(description="Percentage of job description keywords found in resume")
    keyword_frequency_score: float = Field(description="Score based on how often keywords appear in the resume")
    section_completion_percentage: float = Field(description="Percentage of key sections present in the resume")
    formatting_readability_score: float = Field(description="Score for overall formatting and readability")
    hard_soft_skills_balance: float = Field(description="Score for balance between hard technical and soft interpersonal skills")
    proximity_score: float = Field(description="Score for phrases appearing in the right order/context")
    total_ats_score: float = Field(description="Final weighted ATS compatibility score out of 100")
    missing_keywords: List[str] = Field(description="Important keywords from job description missing in resume")
    improvement_suggestions: List[str] = Field(description="Specific suggestions to improve ATS compatibility")
    searchability_suggestions: List[str] = Field(description="Suggestions to improve keyword searchability")
    skills_suggestions: List[str] = Field(description="Suggestions to improve skills presentation")
    formatting_suggestions: List[str] = Field(description="Suggestions to improve formatting and readability")
    section_suggestions: List[str] = Field(description="Suggestions to improve section completeness")
    synonym_suggestions: List[str] = Field(description="Suggestions to improve keyword variation coverage")
    searchability_issues_count: int = Field(description="Number of searchability issues to fix")
    skills_issues_count: int = Field(description="Number of skills‐presentation issues to fix")
    formatting_issues_count: int = Field(description="Number of formatting/readability issues to fix")
    section_issues_count: int = Field(description="Number of section completeness issues to fix")
    synonym_issues_count: int = Field(description="Number of keyword‐variation issues to fix")

class ResumeOptimization(BaseModel):
    improved_summary: str = Field(description="AI-enhanced professional summary")
    improved_bullets: Dict[str, List[str]] = Field(description="Improved bullet points for each experience section")
    suggested_skills: List[str] = Field(description="Additional skills to highlight based on job description")
    formatting_suggestions: List[str] = Field(description="Suggestions for better formatting")
    improved_resume_text: str = Field(description="Full improved resume text")

class CoverLetterOutput(BaseModel):
    cover_letter_text: str = Field(description="Complete cover letter text")

# File processing functions
async def extract_text_from_pdf(file_path):
    return await _sync_extract_text_from_pdf(file_path)

async def _sync_extract_text_from_pdf(file_path):
    # Keep this as async but remove the thread pool executor
    text = ""
    with open(file_path, 'rb') as file:
        pdf_reader = PyPDF2.PdfReader(file)
        for page in pdf_reader.pages:
            text += page.extract_text()
    return text

# Similarly update other extraction functions
async def extract_text_from_docx(file_path):
    return await _sync_extract_text_from_docx(file_path)

async def _sync_extract_text_from_docx(file_path):
    doc = docx.Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

# Simplify the file processing
async def process_resume_file(file, upload_folder):
    filename = secure_filename(file.filename)
    file_path = os.path.join(upload_folder, filename)
    
    # Save file
    with open(file_path, 'wb') as f:
        f.write(file.read())
    
    # Extract text based on file type
    if filename.endswith('.pdf'):
        resume_text = await extract_text_from_pdf(file_path)
    elif filename.endswith('.docx') or filename.endswith('.doc'):
        resume_text = await extract_text_from_docx(file_path)
    else:
        return None, None
    
    return resume_text, file_path

# Chain creation functions
def create_ats_analysis_chain(api_key):
    ats_template = """
    You are an expert ATS (Applicant Tracking System) analyzer and resume optimization specialist.

    I need you to analyze a resume against a specific job description and provide a detailed ATS compatibility score.

    RESUME TEXT:
    {resume_text}

    JOB DESCRIPTION:
    {job_description}

   Analyze the following resume against the provided job description and generate a detailed ATS compatibility report including:

    1. Keyword Match %:
    - Extract ALL critical keywords and key phrases (skills, certifications, tools, industry jargon, job titles).
    - Assess exact and semantic keyword matches, including synonyms and relevant variations.
    - Evaluate keyword placement priority: prioritize headline, professional summary, skills section, and work experience.
    - Score keyword density ensuring natural language flow (avoid stuffing).
    - Identify missing high-impact keywords critical to the job role and recommend optimal placement.

    2. Section Completion %:
    - Verify presence of core ATS sections with professional formatting.
    - Score based on alignment with job requirements and ATS parsing ease.
    - Recommend missing or reorder sections for maximum ATS compatibility.

    3. Formatting and Readability Score:
    Assess resume formatting and readability for ATS systems, including:
    - Use of standard fonts (e.g., Arial, Calibri)
    - Consistent date formats (e.g., MM/YYYY)
    - Absence of graphics, tables, columns, headers, or footers
    - Clear section headings and logical structure
    - Appropriate file format (.docx or .pdf)

    4. Hard vs Soft Skills Balance:
    Evaluate the balance between technical (hard) and interpersonal (soft) skills aligned to the job description.

    5. Proximity Score:
    Analyze whether related keywords and phrases appear close together and in the correct context.

    6. Final ATS Score:
    Calculate a weighted ATS score out of 100 using the formula:
    (Keyword Match * 0.35) + (Section Completion * 0.25) + (Formatting * 0.20) + (Skills Balance * 0.10) + (Proximity * 0.10)

    7. List important keywords from the job description missing in the resume

    Additionally, organize your suggestions into the following categories to match with the job description and get a better ATS score:
    1. Searchability suggestions - How to improve keyword match and searchability
       • Also provide `searchability_issues_count`: the number of searchability issues to fix
    2. Skills suggestions - How to improve hard/soft skills balance and presentation
       • Also provide `skills_issues_count`: the number of skills-presentation issues to fix
    3. Formatting suggestions - How to improve layout and ATS readability
       • Also provide `formatting_issues_count`: the number of formatting/readability issues to fix
    4. Section suggestions - How to improve section completeness
       • Also provide `section_issues_count`: the number of section completeness issues to fix
    5. Synonym suggestions - How to improve keyword variation coverage
       • Also provide `synonym_issues_count`: the number of keyword-variation issues to fix

    For each category:
    1. List all individual issues that need fixing.
    2. Provide 2–3 specific, actionable suggestions.
    3. Emit the corresponding `<category>_issues_count` integer, **exactly equal** to the number of issues you listed above.

    {format_instructions}
    """
    
    llm = ChatOpenAI(
        model_name="gpt-4o-mini",
        temperature=0.2,
        openai_api_key=api_key
    )
    
    parser = PydanticOutputParser(pydantic_object=ATSScore)
    prompt = PromptTemplate(
        template=ats_template,
        input_variables=["resume_text", "job_description"],
        partial_variables={"format_instructions": parser.get_format_instructions()}
    )

    return RunnableSequence(prompt, llm, parser)

def create_resume_optimization_chain(api_key):
    resume_template = """
    You are an expert resume writer, career coach, and Applicant Tracking System (ATS) specialist. Your mission is to transform the candidate’s existing resume into a highly optimized, keyword‑rich document that perfectly aligns with the given job description—while **preserving every original section** of the resume.

    **INPUTS:**  
    - **RESUME TEXT:** `{resume_text}`  
    - **JOB DESCRIPTION:** `{job_description}`  
    - **ATS ANALYSIS:** `{ats_analysis}`  

    ---

    ## Instructions

    1. **Section Preservation & Enhancement**  
    - **Retain every section** from the original resume (e.g., Professional Summary, Experience, Projects, Education, Skills, Certifications, Volunteer, etc.).  
    - For each section, **refine content** to mirror the job description’s language and keywords.

    2. **Skill Extraction & Mapping**  
    - Parse the JD for all **required** and **preferred** skills/technologies/methodologies.  
    - Build a “Skills Checklist” and ensure each skill appears in the resume—either in the Skills section or embedded within relevant bullets.  
    - If the candidate legitimately holds a JD‑listed skill that’s missing, add it; do not invent.

    3. **Professional Summary (1–3 sentences)**  
    - Incorporate the top 4–6 JD keywords/phrases.  
    - Highlight years of experience, core qualifications, and career goal.  
    - Keep it concise, impactful, and natural.

    4. **Experience Section Optimization**  
    - Maintain **reverse‑chronological** order and use present‑tense for current roles, past‑tense for prior roles.  
    - **Rewrite each bullet** to:  
        - Use strong action verbs (e.g., “Led,” “Architected,” “Streamlined”).  
        - Embed exact JD keywords (e.g., “CI/CD pipelines,” “Agile,” “AWS”).  
        - Quantify achievements (e.g., “Reduced deployment time by 40%,” “Managed a $200K budget”).  
    - **Retain all original positions** and headers; enhance every bullet for relevance and metrics.

    5. **Skills Section Enhancement**  
    - Under a clear **Skills** heading, list every core JD skill/term as individual bullets, ordered by relevance.  
    - If a required skill is absent in the original, note the gap and suggest a related, transferable skill.

    6. **Education & Certifications**  
    - **Preserve** all original entries.  
    - Standardize formatting: *Degree/Certification* – *Institution/Issuer*, *Month Year*.  
    - Add any JD‑required certification the candidate holds or is pursuing.

    7. **Formatting for Optimal ATS Parsing**  
    - Single‑column, reverse‑chronological layout with standard headings.  
    - Simple bullets (– or ●); no tables, graphics, or columns.  
    - Standard font (Arial or Calibri, 10–12 pt); consistent spacing, margins, and date format (e.g., “MMM YYYY – MMM YYYY”).  
    - Recommend final filename: `Firstname_Lastname_Resume.pdf`.

    8. **Full Resume Rewrite**  
    - Provide a polished 1–2 page document that:  
        1. **Professional Summary** (refined)  
        2. **Experience** (all original roles, optimized)  
        3. **Skills** (exhaustive, JD‑aligned)  
        4. **Education** (standardized)  
        5. **Certifications** (standardized)  
        6. **Other Sections** (e.g., Projects, Volunteer) – refined to include relevant keywords.  
    - Ensure **every JD keyword** appears at least once across summary, skills, or experience—without unnatural repetition.

    9. **Final Verification**  
    - Confirm 100% coverage of required/preferred JD skills.  
    - Ensure natural keyword density and no keyword stuffing.  
    - Validate that each original section is present and improved.  
    - Include formatting notes in curly braces (e.g., `{format_instructions}`) for the user’s final layout.

    **OUTPUT:**  
    A complete ATS‑optimized resume—retaining and enhancing all original sections.
    """
    
    llm = ChatOpenAI(
        model_name="gpt-4o-mini",
        temperature=0.2,
        openai_api_key=api_key
    )
    
    parser = PydanticOutputParser(pydantic_object=ResumeOptimization)
    prompt = PromptTemplate(
        template=resume_template,
        input_variables=["resume_text", "job_description", "ats_analysis"],
        partial_variables={"format_instructions": parser.get_format_instructions()}
    )

    return RunnableSequence(prompt, llm, parser)

def create_cover_letter_chain(api_key):
    current_date = datetime.now().strftime("%B %d, %Y")
    cover_letter_template = """
    You are an expert cover letter writer with deep knowledge of professional communication and hiring practices.
    
    Create a compelling, personalized cover letter based on the resume and job description provided.
    
    RESUME TEXT:
    {resume_text}
    
    JOB DESCRIPTION:
    {job_description}

    CURRENT DATE: {current_date}
    
    Please write a professional cover letter that:
    1. Has a proper business letter format with date and contact information
    2. Opens with a compelling introduction that shows enthusiasm for the position
    3. Highlights 2-3 key qualifications from the resume that directly match the job requirements
    4. Incorporates important keywords from the job description naturally
    5. Explains why the applicant is a good fit for the company and role specifically
    6. Closes with a strong call to action and thank you
    7. Maintains a professional yet personable tone
    8. Is approximately 250-350 words in total
    
    {format_instructions}
    """
    
    llm = ChatOpenAI(
        model_name="gpt-4o-mini",
        temperature=0.2,
        openai_api_key=api_key
    )
    
    parser = PydanticOutputParser(pydantic_object=CoverLetterOutput)
    prompt = PromptTemplate(
        template=cover_letter_template,
        input_variables=["resume_text", "job_description"],
        partial_variables={
            "format_instructions": parser.get_format_instructions(),
            "current_date": current_date
        }
    )
    
    return RunnableSequence(prompt, llm, parser)


# Document creation functions
async def create_docx_document(content, document_type="resume"):
    return await _sync_create_docx_document(content, document_type)

async def _sync_create_docx_document(content, document_type):
    doc = DocxDocument()
    sections = content.split('\n\n')
    for section in sections:
        if section.strip():
            doc.add_paragraph(section)
    file_obj = BytesIO()
    doc.save(file_obj)
    file_obj.seek(0)
    return file_obj

async def create_pdf_document(content, document_type="resume"):
    return await _sync_create_pdf_document(content, document_type)

async def _sync_create_pdf_document(content, document_type):
    file_obj = BytesIO()
    doc = SimpleDocTemplate(file_obj, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    normal_style = styles['Normal']
    heading_style = styles['Heading1']
    
    sections = content.split('\n\n')
    for section in sections:
        if section.strip():
            if len(section) < 50 and section.isupper() or ':' in section and len(section.split(':')[0]) < 20:
                story.append(Paragraph(section, heading_style))
            else:
                story.append(Paragraph(section, normal_style))
            story.append(Spacer(1, 12))
    
    doc.build(story)
    file_obj.seek(0)
    return file_obj